"""
docx_pipeline_router.py
───────────────────────
Self-contained FastAPI router — drop this single file into any FastAPI project.

USAGE in your existing main.py:
    from docx_pipeline_router import router as docx_router
    app.include_router(docx_router, prefix="/docx", tags=["DOCX Pipeline"])

ENDPOINTS (all under whatever prefix you choose):
    POST /upload              → upload .doc/.docx, returns job_id
    GET  /status/{job_id}     → poll processing status
    GET  /download/{job_id}/markdown   → download .md file
    GET  /download/{job_id}/semantic   → download semantic .json file

OUTPUT FOLDER:
    uploads/<file_stem>/
        <file_stem>.md
        <file_stem>.dom.json
        <file_stem>.semantic.json
        images/
            picture_0.png
            ...

REQUIRED ENV VARS (see .env.docx_pipeline):
    AWS_REGION, AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY
    BEDROCK_MODEL_ID, USE_VLM, USE_TEXTRACT
    DOCX_OUTPUT_DIR   (default: ./uploads)
"""

from __future__ import annotations

import io
import json
import logging
import os
import shutil
import uuid
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from typing import Any, Optional

import boto3
from botocore.exceptions import BotoCoreError, ClientError
from dotenv import load_dotenv
from fastapi import APIRouter, BackgroundTasks, File, HTTPException, UploadFile
from fastapi.responses import FileResponse
from PIL import Image
from pydantic import BaseModel

load_dotenv()

# ── Logging ────────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
)
log = logging.getLogger("docx_pipeline")

# ── Config from environment ────────────────────────────────────────────────────
AWS_REGION             = os.getenv("AWS_REGION", "us-east-1")
AWS_ACCESS_KEY_ID      = os.getenv("AWS_ACCESS_KEY_ID", "")
AWS_SECRET_ACCESS_KEY  = os.getenv("AWS_SECRET_ACCESS_KEY", "")
BEDROCK_MODEL_ID       = os.getenv("BEDROCK_MODEL_ID", "amazon.nova-lite-v1:0")
USE_VLM                = os.getenv("USE_VLM", "true").lower() == "true"
USE_TEXTRACT           = os.getenv("USE_TEXTRACT", "true").lower() == "true"
OUTPUT_DIR             = Path(os.getenv("DOCX_OUTPUT_DIR", "./uploads"))
IMAGES_SCALE           = 2.0
TEXTRACT_MAX_BYTES     = 5 * 1024 * 1024
MIN_CHART_AREA         = 40_000
MIN_OCR_TOKENS         = 2
MAX_OCR_TOKENS         = 80
ALLOWED_EXT            = {".doc", ".docx"}
STAGING_DIR            = OUTPUT_DIR / "_staging"

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
STAGING_DIR.mkdir(parents=True, exist_ok=True)

# ── In-memory job store ────────────────────────────────────────────────────────
_jobs: dict[str, dict[str, Any]] = {}
_executor = ThreadPoolExecutor(max_workers=2)

# ── Pydantic schemas ───────────────────────────────────────────────────────────
class UploadResponse(BaseModel):
    job_id: str
    status: str
    message: str

class JobStatus(BaseModel):
    job_id: str
    status: str
    document: Optional[str] = None
    tables_count: Optional[int] = None
    images_count: Optional[int] = None
    markdown_preview: Optional[str] = None
    markdown_path: Optional[str] = None
    semantic_json_path: Optional[str] = None
    dom_json_path: Optional[str] = None
    error: Optional[str] = None

# ── Router ─────────────────────────────────────────────────────────────────────
router = APIRouter()

# ══════════════════════════════════════════════════════════════════════════════
# AWS HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _aws_kwargs() -> dict:
    kw: dict[str, Any] = {"region_name": AWS_REGION}
    if AWS_ACCESS_KEY_ID and AWS_SECRET_ACCESS_KEY:
        kw["aws_access_key_id"] = AWS_ACCESS_KEY_ID
        kw["aws_secret_access_key"] = AWS_SECRET_ACCESS_KEY
    return kw


# ══════════════════════════════════════════════════════════════════════════════
# OCR — Amazon Textract
# ══════════════════════════════════════════════════════════════════════════════

def _img_to_bytes_for_textract(img: Image.Image) -> bytes:
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    data = buf.getvalue()
    if len(data) <= TEXTRACT_MAX_BYTES:
        return data
    for q in (90, 75, 60, 45):
        buf = io.BytesIO()
        img.convert("RGB").save(buf, format="JPEG", quality=q)
        data = buf.getvalue()
        if len(data) <= TEXTRACT_MAX_BYTES:
            return data
    return data


def run_ocr(img: Image.Image) -> str:
    if not USE_TEXTRACT:
        return ""
    try:
        image_bytes = _img_to_bytes_for_textract(img)
        client = boto3.client("textract", **_aws_kwargs())
        response = client.detect_document_text(Document={"Bytes": image_bytes})
        lines = [
            b["Text"] for b in response.get("Blocks", [])
            if b.get("BlockType") == "LINE"
        ]
        return "\n".join(lines).strip()
    except (BotoCoreError, ClientError) as exc:
        log.warning("Textract OCR failed: %s", exc)
        return ""
    except Exception as exc:
        log.warning("OCR error: %s", exc)
        return ""


# ══════════════════════════════════════════════════════════════════════════════
# CHART DETECTION — AWS Bedrock Nova Lite
# ══════════════════════════════════════════════════════════════════════════════

_CHART_PROMPT = (
    "You are analyzing an image extracted from a business document. "
    "Determine whether the image is a chart, graph, diagram, or table of data. "
    "Return ONLY valid JSON — no prose, no markdown fences — matching this shape:\n"
    '{"chart_type": "<bar|line|pie|table|diagram|not_a_chart>", '
    '"series": [{"label": "<string>", "value": "<string>"}], '
    '"summary": "<one sentence describing the chart>"}\n'
    "If the image is a photo, logo, signature, or decorative element, "
    'return {"chart_type": "not_a_chart", "series": [], "summary": ""}.'
)

_NOT_EVALUATED = {"chart_type": "not_evaluated", "series": [], "summary": ""}
_ERROR_RESULT   = {"chart_type": "error",         "series": [], "summary": ""}


def looks_like_chart(img: Image.Image, ocr_text: str) -> bool:
    area = img.width * img.height
    tokens = len(ocr_text.split())
    return area > MIN_CHART_AREA and MIN_OCR_TOKENS <= tokens <= MAX_OCR_TOKENS


def describe_chart(image_path: Path) -> dict[str, Any]:
    try:
        with Image.open(image_path) as im:
            buf = io.BytesIO()
            im.save(buf, format="PNG")
            image_bytes = buf.getvalue()

        client = boto3.client("bedrock-runtime", **_aws_kwargs())
        response = client.converse(
            modelId=BEDROCK_MODEL_ID,
            messages=[{
                "role": "user",
                "content": [
                    {"image": {"format": "png", "source": {"bytes": image_bytes}}},
                    {"text": _CHART_PROMPT},
                ],
            }],
            inferenceConfig={"maxTokens": 512, "temperature": 0.1},
        )
        raw = response["output"]["message"]["content"][0]["text"].strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
            raw = raw.strip()
        parsed = json.loads(raw)
        parsed.setdefault("chart_type", "unknown")
        parsed.setdefault("series", [])
        parsed.setdefault("summary", "")
        return parsed
    except (BotoCoreError, ClientError) as exc:
        log.error("Bedrock call failed: %s", exc)
        return {**_ERROR_RESULT, "summary": str(exc)}
    except json.JSONDecodeError as exc:
        log.warning("Bedrock non-JSON response: %s", exc)
        return {**_ERROR_RESULT, "summary": f"JSON parse error: {exc}"}
    except Exception as exc:
        log.warning("VLM error: %s", exc)
        return {**_ERROR_RESULT, "summary": str(exc)}


# ══════════════════════════════════════════════════════════════════════════════
# TABLE EXTRACTOR
# ══════════════════════════════════════════════════════════════════════════════

def _safe(val: Any) -> Any:
    if val is None or isinstance(val, (str, int, float, bool)):
        return val
    return str(val)


def extract_tables(doc) -> list[dict[str, Any]]:
    out = []
    for i, table in enumerate(doc.tables):
        records = None
        try:
            try:
                df = table.export_to_dataframe(doc)
            except TypeError:
                df = table.export_to_dataframe()
            records = [{k: _safe(v) for k, v in row.items()} for row in df.to_dict(orient="records")]
        except Exception as exc:
            log.warning("Table %d dataframe export failed: %s", i, exc)

        raw_cells: list[dict] = []
        try:
            for cell in table.data.table_cells:
                raw_cells.append({
                    "text":     cell.text,
                    "row":      cell.start_row_offset_idx,
                    "col":      cell.start_col_offset_idx,
                    "row_span": cell.row_span,
                    "col_span": cell.col_span,
                    "is_header": getattr(cell, "column_header", False),
                })
        except Exception as exc:
            log.warning("Table %d raw cell export failed: %s", i, exc)

        caption = None
        try:
            if getattr(table, "captions", None):
                caption = table.caption_text(doc)
        except Exception:
            pass

        prov = table.prov[0] if getattr(table, "prov", None) else None
        out.append({
            "table_index":       i,
            "caption":           caption,
            "page":              getattr(prov, "page_no", None),
            "num_rows":          table.data.num_rows if table.data else None,
            "num_cols":          table.data.num_cols if table.data else None,
            "flattened_records": records,
            "raw_cells":         raw_cells,
        })
    return out


# ══════════════════════════════════════════════════════════════════════════════
# IMAGE EXTRACTOR
# ══════════════════════════════════════════════════════════════════════════════

def _get_pil_image(picture, doc) -> Optional[Image.Image]:
    try:
        img = picture.get_image(doc)
        if img is not None:
            return img
    except Exception:
        pass
    try:
        img = getattr(picture, "image", None)
        if isinstance(img, Image.Image):
            return img
        pil = getattr(img, "pil_image", None)
        if pil is not None:
            return pil
    except Exception:
        pass
    return None


def extract_images(doc, image_dir: Path) -> list[dict[str, Any]]:
    image_dir.mkdir(parents=True, exist_ok=True)
    out = []
    for i, picture in enumerate(doc.pictures):
        img = _get_pil_image(picture, doc)
        if img is None:
            log.info("Picture %d: no image data — skipping.", i)
            continue

        img_path = image_dir / f"picture_{i}.png"
        try:
            img.save(img_path)
        except Exception as exc:
            log.warning("Could not save picture %d: %s", i, exc)
            continue

        caption = None
        try:
            if getattr(picture, "captions", None):
                caption = picture.caption_text(doc)
        except Exception:
            pass

        prov = picture.prov[0] if getattr(picture, "prov", None) else None
        bbox_obj = getattr(prov, "bbox", None)
        bbox = ({"l": bbox_obj.l, "t": bbox_obj.t, "r": bbox_obj.r, "b": bbox_obj.b}
                if bbox_obj else None)

        ocr_text = run_ocr(img)

        if USE_VLM and looks_like_chart(img, ocr_text):
            log.info("Picture %d looks chart-like — calling Bedrock", i)
            semantic = describe_chart(img_path)
        else:
            semantic = _NOT_EVALUATED.copy()

        out.append({
            "picture_index":     i,
            "file":              str(img_path),
            "caption":           caption,
            "page":              getattr(prov, "page_no", None),
            "bbox":              bbox,
            "width":             img.width,
            "height":            img.height,
            "ocr_text":          ocr_text,
            "semantic":          semantic,
            "resolution_method": "docling_dom",
        })
    return out


# ══════════════════════════════════════════════════════════════════════════════
# MARKDOWN RENDERER
# ══════════════════════════════════════════════════════════════════════════════

import re as _re


def _fix_table_cell_bullets(md: str) -> str:
    """
    Docling flattens intra-cell bullet lists from Word tables into a single
    string like '- Ooo - Abc - Efgh'.
    Detects cells that START with '- ' and contain more ' - ' separators,
    and replaces them with '<br>- ' so bullets render correctly.
    """
    def _fix_cell(cell: str) -> str:
        c = cell.strip()
        if _re.match(r'^-\s+\S', c) and ' - ' in c:
            c = _re.sub(r'\s+-\s+', '<br>- ', c)
        return c

    fixed_lines = []
    for line in md.split('\n'):
        if line.startswith('|') and not _re.match(r'^\|[-| :]+\|$', line.strip()):
            parts = line.split('|')
            fixed_parts = [
                _fix_cell(p) if 0 < i < len(parts) - 1 else p
                for i, p in enumerate(parts)
            ]
            line = '|'.join(fixed_parts)
        fixed_lines.append(line)
    return '\n'.join(fixed_lines)


def _extract_nested_table_cells(docx_path: Path) -> dict[str, str]:
    """
    Use python-docx to find outer table cells whose content should be
    reformatted as structured data. Two cases are handled:

    Case A — Nested Word tables:
      A cell contains an actual inner <w:tbl> table.

    Case B — Manually typed tabular data:
      A cell contains multiple paragraphs where each non-blank paragraph
      has 2+ whitespace-separated short tokens that look like data rows.
      E.g. user typed:  "A simple table" [Enter] "A  10" [Enter] "B  15" [Enter] "C  20"
      Docling flattens: "A simple table A 10 B 15 C 20"
      We restore:       "A simple table<br>A \\| 10<br>B \\| 15<br>C \\| 20"

    Returns {flat_key: formatted_replacement} for use in _apply_nested_table_fixes.
    """
    try:
        from docx import Document as _DocxDocument
        doc = _DocxDocument(str(docx_path))
    except Exception as exc:
        log.warning("python-docx could not open '%s': %s", docx_path.name, exc)
        return {}

    replacements: dict[str, str] = {}

    def _looks_like_data_row(text: str) -> list[str] | None:
        parts = _re.split(r'[ \t]{2,}|\t', text.strip())
        parts = [p.strip() for p in parts if p.strip()]
        if len(parts) >= 2 and all(len(p) <= 40 for p in parts):
            return parts
        return None

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:

                # ── Case A: cell contains a nested Word table ──────────────
                if cell.tables:
                    para_text = ' '.join(
                        p.text.strip() for p in cell.paragraphs if p.text.strip()
                    )
                    nested_rows: list[str] = []
                    for nested_table in cell.tables:
                        for nested_row in nested_table.rows:
                            cols = [
                                c.text.strip()
                                for c in nested_row.cells
                                if c.text.strip()
                            ]
                            if cols:
                                nested_rows.append(' \\| '.join(cols))

                    parts_list: list[str] = []
                    if para_text:
                        parts_list.append(para_text)
                    parts_list.extend(nested_rows)
                    formatted = '<br>'.join(parts_list)

                    all_nested_text = ' '.join(
                        c.text.strip()
                        for nt in cell.tables
                        for nr in nt.rows
                        for c in nr.cells
                        if c.text.strip()
                    )
                    flat_raw = ' '.join(filter(None, [para_text, all_nested_text]))
                    flat_key = ' '.join(flat_raw.split())

                    if flat_key and formatted and flat_key != formatted:
                        replacements[flat_key] = formatted
                    continue

                # ── Case B: manually typed tabular data ────────────────────
                non_blank = [p for p in cell.paragraphs if p.text.strip()]
                if len(non_blank) < 2:
                    continue

                classified: list[tuple] = []
                for p in non_blank:
                    tokens = _looks_like_data_row(p.text)
                    if tokens:
                        classified.append(('row', tokens))
                    else:
                        classified.append(('text', p.text.strip()))

                row_count = sum(1 for kind, _ in classified if kind == 'row')
                if row_count < 2:
                    continue

                fmt_parts: list[str] = []
                for kind, content in classified:
                    if kind == 'text':
                        fmt_parts.append(content)
                    else:
                        fmt_parts.append(' \\| '.join(content))
                formatted = '<br>'.join(fmt_parts)

                flat_key = ' '.join(
                    ' '.join(p.text.split()) for p in non_blank
                )
                flat_key = ' '.join(flat_key.split())

                if flat_key and formatted and flat_key != formatted:
                    replacements[flat_key] = formatted

    return replacements


def _apply_nested_table_fixes(md: str, fixes: dict[str, str]) -> str:
    """Replace flattened nested-table text inside pipe-table cells."""
    if not fixes:
        return md

    fixed_lines = []
    for line in md.split('\n'):
        if line.startswith('|') and not _re.match(r'^\|[-| :]+\|$', line.strip()):
            for flat, formatted in fixes.items():
                # Match the flat text inside a cell (between | delimiters),
                # normalising internal whitespace before comparing.
                normalised_line = ' '.join(line.split())
                if flat in normalised_line:
                    # Rebuild the line using the original but replacing the flat text
                    # We normalise cell content for matching, then put formatted back.
                    cells = line.split('|')
                    new_cells = []
                    for ci, cell in enumerate(cells):
                        cell_norm = ' '.join(cell.split())
                        if flat in cell_norm:
                            new_cells.append(' ' + formatted + ' ')
                        else:
                            new_cells.append(cell)
                    line = '|'.join(new_cells)
        fixed_lines.append(line)
    return '\n'.join(fixed_lines)


def _fix_html_entities(md: str) -> str:
    """
    Docling HTML-encodes special characters in headings and table cells.
    Decode the most common ones back to readable text.
    """
    return (
        md
        .replace("&amp;",  "&")
        .replace("&lt;",   "<")
        .replace("&gt;",   ">")
        .replace("&quot;", '"')
        .replace("&#39;",  "'")
        .replace("&nbsp;", " ")
    )


def _escape_block_html_in_lists(md: str) -> str:
    """
    Markdown renderers treat block-level HTML tags (<ol>, </ol>, <ul>, etc.)
    as real HTML even inside list item text.  When a list item contains text
    like "(e.g., <ul>/<ol> nesting in HTML)", the renderer interprets </ol>
    as the closing tag of the current ordered list, breaking the numbering.

    This escapes those tags to HTML entities inside list item lines ONLY,
    so they render as visible text.  Our own <br> tags (used in table cells
    on pipe-table lines) are unaffected because they're never on list lines.
    """
    # Block-level tags that break list rendering when left raw
    _BLOCK_TAGS = _re.compile(
        r'<(/?(?:ul|ol|li|dl|dt|dd|table|thead|tbody|tfoot|tr|th|td'
        r'|div|section|article|aside|header|footer|main|nav|p'
        r'|blockquote|pre|figure|figcaption)\b[^>]*)>',
        _re.IGNORECASE,
    )
    list_line_re = _re.compile(r'^(\s*(?:[-*+]|\d+\.)\s+)(.+)$')

    out = []
    for line in md.split('\n'):
        m = list_line_re.match(line)
        if m:
            prefix, content = m.group(1), m.group(2)
            content = _BLOCK_TAGS.sub(lambda mo: f'&lt;{mo.group(1)}&gt;', content)
            line = prefix + content
        out.append(line)
    return '\n'.join(out)


def _fix_merged_cell_duplicates(md: str) -> str:
    """
    Docling represents colspan merged cells by repeating the cell text,
    e.g.:  | **Region & Quarter**  **Region & Quarter** ||
    and rowspan gaps as empty cells:  ||

    This cleans up:
      - Cells where the same text (or bold text) is repeated twice → keep once
      - Trailing empty phantom cells from colspan (the extra ||)
      - Empty rowspan cells → replace with ↑ (indicates 'merged from above')
    """
    def _dedup_cell(cell: str) -> str:
        c = cell.strip()
        if not c:
            return cell  # handled separately below

        # Pattern: "TEXT  TEXT" or "**TEXT**  **TEXT**" (same content repeated)
        # Split on 2+ spaces and check if all parts are the same
        parts = _re.split(r'\s{2,}', c)
        if len(parts) >= 2:
            unique = list(dict.fromkeys(p.strip() for p in parts if p.strip()))
            if len(unique) == 1:
                return ' ' + unique[0] + ' '

        return cell

    fixed_lines = []
    for line in md.split('\n'):
        if line.startswith('|') and not _re.match(r'^\|[-| :]+\|$', line.strip()):
            cells = line.split('|')
            new_cells = []
            for i, cell in enumerate(cells):
                if i == 0 or i == len(cells) - 1:
                    new_cells.append(cell)
                    continue
                if cell.strip() == '':
                    # Empty cell — could be colspan phantom or rowspan gap
                    # Replace with a subtle marker only if surrounded by real cells
                    new_cells.append(' ↑ ')
                else:
                    new_cells.append(_dedup_cell(cell))
            line = '|'.join(new_cells)
        fixed_lines.append(line)
    return '\n'.join(fixed_lines)


def _fix_list_blank_lines(md: str) -> str:
    """
    Docling adds blank lines between every list item, which breaks Markdown
    list nesting. A blank line between two list items tells the parser to
    end the current list and start a new one — destroying hierarchy.

    This removes blank lines that appear BETWEEN consecutive list items
    (lines starting with optional spaces + '- '/'* '/'+ ' '/'N. ').

    Blank lines before/after non-list content are left untouched.
    """
    list_re = _re.compile(r'^\s*([-*+]|\d+\.)\s+')

    lines = md.split('\n')
    out: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip() == '':
            # Look backwards (skip multiple consecutive blank lines)
            prev_idx = len(out) - 1
            while prev_idx >= 0 and out[prev_idx].strip() == '':
                prev_idx -= 1
            prev_is_list = prev_idx >= 0 and bool(list_re.match(out[prev_idx]))

            # Look forwards (skip multiple consecutive blank lines)
            next_idx = i + 1
            while next_idx < len(lines) and lines[next_idx].strip() == '':
                next_idx += 1
            next_is_list = next_idx < len(lines) and bool(list_re.match(lines[next_idx]))

            if prev_is_list and next_is_list:
                # Drop this blank line (and skip any further consecutive blanks)
                i = next_idx
                continue
        out.append(line)
        i += 1
    return '\n'.join(out)


def _fix_mixed_list_nesting(md: str) -> str:
    """
    Fixes two related Docling rendering problems for mixed-type nested lists:

    Problem A — Bullets not indented under numbered items:
      Docling renders:
          5. Tables
          - Simple Tables      ← should be '   - Simple Tables'
          - Complex Tables
          1. Columns           ← restarts at 1, should be 6
      Expected:
          5. Tables
             - Simple Tables
             - Complex Tables
          6. Columns

    Algorithm:
      - Walk lines. When a numbered item (N. text) is seen, record N.
      - If the very next line (no blank) is an unindented bullet (- text):
          → enter "sub-bullet mode": indent all following unindented bullets
      - When sub-bullet mode ends and the next numbered item is "1. text":
          → renumber it to (last_num + 1) and continue renumbering.
    """
    numbered_re = _re.compile(r'^(\d+)\.\s+(.+)$')
    bullet_re   = _re.compile(r'^([-*+])\s+')

    lines = md.split('\n')
    out: list[str] = []
    last_num = 0          # last seen numbered list counter
    in_sub = False        # currently inside a sub-bullet block under a numbered item
    sub_continuation = 0  # what the next numbered item should be after sub-bullets

    for i, line in enumerate(lines):
        nm = numbered_re.match(line)
        bm = bullet_re.match(line)

        if nm:
            num = int(nm.group(1))
            if in_sub:
                # Numbered list resuming after sub-bullets
                in_sub = False
                if num == 1 and sub_continuation > 0:
                    # Restart detected — renumber to continue sequence
                    line = f'{sub_continuation}. {nm.group(2)}'
                    last_num = sub_continuation
                    sub_continuation = 0
                else:
                    last_num = num
            else:
                last_num = num

            out.append(line)

        elif bm and not line.startswith(' ') and not line.startswith('\t'):
            # Unindented bullet
            if not in_sub:
                # Check if the previous non-empty output line was a numbered item
                prev = next((l for l in reversed(out) if l.strip()), '')
                if numbered_re.match(prev):
                    in_sub = True
                    sub_continuation = last_num + 1

            if in_sub:
                out.append('   ' + line)   # indent 3 spaces = sub-item
            else:
                out.append(line)

        else:
            if line.strip() and not bullet_re.match(line.lstrip()):
                # Non-list, non-blank content ends sub-bullet mode
                in_sub = False
            out.append(line)

    return '\n'.join(out)


def _build_list_level_map(docx_path: Path) -> dict[str, int]:
    """
    Open the DOCX with python-docx and determine each list paragraph's
    indent level.  Returns {normalised_text: level} (level 0 = top level).

    Three-tier lookup in order of priority:
      1. Direct w:numPr/w:ilvl on the paragraph element
      2. Style name pattern  (e.g. "List Bullet 2" → level 1,
                                   "List Number 3"  → level 2)
      3. w:numPr inherited from the paragraph style chain
    """
    # Known Word list-style-name → level mappings
    _STYLE_LEVEL_NAMES = {
        "list bullet":    0, "list bullet 2": 1, "list bullet 3": 2,
        "list bullet 4":  3, "list bullet 5": 4,
        "list number":    0, "list number 2": 1, "list number 3": 2,
        "list number 4":  3, "list number 5": 4,
        "list paragraph": 0, "list continue": 0,
        "list continue 2": 1, "list continue 3": 2,
    }
    import re as _re2

    try:
        from docx import Document as _DocxDocument
        from docx.oxml.ns import qn as _qn
        doc = _DocxDocument(str(docx_path))
        level_map: dict[str, int] = {}

        def _ilvl_from_numPr(el) -> Optional[int]:
            """Extract ilvl value from a w:numPr element."""
            ilvl_el = el.find(_qn('w:ilvl'))
            if ilvl_el is not None:
                return int(ilvl_el.get(_qn('w:val'), 0))
            return None

        def _get_level(para) -> Optional[int]:
            # --- Tier 1: direct w:numPr on the paragraph ---
            pPr = para._p.find(_qn('w:pPr'))
            if pPr is not None:
                numPr = pPr.find(_qn('w:numPr'))
                if numPr is not None:
                    lv = _ilvl_from_numPr(numPr)
                    if lv is not None:
                        return lv

            # --- Tier 2: style name encodes the level ---
            if para.style:
                sname = para.style.name.strip().lower()
                if sname in _STYLE_LEVEL_NAMES:
                    return _STYLE_LEVEL_NAMES[sname]
                # Catch styles like "List Bullet2" or "Bullet 3" etc.
                m = _re2.search(r'\b(\d)\s*$', sname)
                if m and any(k.rstrip(' 0123456789') in sname
                             for k in ("list bullet", "list number",
                                       "list paragraph", "list continue")):
                    return int(m.group(1)) - 1

            # --- Tier 3: inherited w:numPr from style chain ---
            try:
                style = para.style
                visited: set = set()
                while style is not None and id(style) not in visited:
                    visited.add(id(style))
                    for pPr_el in style.element.iter(_qn('w:pPr')):
                        numPr = pPr_el.find(_qn('w:numPr'))
                        if numPr is not None:
                            lv = _ilvl_from_numPr(numPr)
                            if lv is not None:
                                return lv
                    style = getattr(style, 'base_style', None)
            except Exception:
                pass

            return None

        for para in doc.paragraphs:
            level = _get_level(para)
            if level is not None:
                text = ' '.join(para.text.split())
                if text:
                    level_map[text] = level

        log.debug("_build_list_level_map: found %d list items", len(level_map))
        return level_map

    except Exception as exc:
        log.warning("_build_list_level_map failed: %s", exc)
        return {}



def _apply_list_indentation(md: str, level_map: dict[str, int]) -> str:
    """
    For each list line in the markdown, look up its text in level_map and
    prepend the correct number of spaces (4 per level).

    Also renumbers top-level ordered list items that restarted at 1 because
    a sub-list broke the sequence.
    """
    if not level_map:
        return md

    item_re = _re.compile(r'^(\s*)([-*+]|\d+\.)\s+(.+)$')

    lines = md.split('\n')
    out: list[str] = []
    # Track ordered list counters per level: {level: counter}
    ol_counters: dict[int, int] = {}

    for line in lines:
        m = item_re.match(line)
        if not m:
            out.append(line)
            continue

        existing_indent = m.group(1)
        marker = m.group(2)
        text = m.group(3)
        norm = ' '.join(text.split())

        level = level_map.get(norm)
        if level is None:
            # Text not found in map — keep line as-is
            out.append(line)
            continue

        indent = '    ' * level  # 4 spaces per level

        # For numbered markers, track and fix the counter per level
        if _re.match(r'^\d+\.$', marker):
            # Reset deeper-level counters when we go up a level
            for l in list(ol_counters.keys()):
                if l > level:
                    del ol_counters[l]
            ol_counters[level] = ol_counters.get(level, 0) + 1
            marker = f'{ol_counters[level]}.'

        out.append(f'{indent}{marker} {text}')

    return '\n'.join(out)


def render_markdown(base_md: str, tables: list, images: list,
                    docx_path: Optional[Path] = None) -> str:
    # 1. Fix flattened bullet lists inside table cells
    base_md = _fix_table_cell_bullets(base_md)

    # 2. Fix cells containing nested tables (mini-tables flattened by Docling)
    if docx_path is not None:
        nested_fixes = _extract_nested_table_cells(docx_path)
        base_md = _apply_nested_table_fixes(base_md, nested_fixes)

    # 3. Decode HTML entities Docling introduces (& → &amp; etc.)
    base_md = _fix_html_entities(base_md)

    # 4. Clean up merged-cell duplicate text in colspan/rowspan headers
    base_md = _fix_merged_cell_duplicates(base_md)

    # 5. Remove blank lines Docling inserts between list items (destroys nesting)
    base_md = _fix_list_blank_lines(base_md)

    # 6. Fix mixed-type list nesting (bullets under numbered, and list restarts)
    base_md = _fix_mixed_list_nesting(base_md)

    # 7. Apply proper indentation from DOCX ilvl levels (fixes nested lists)
    if docx_path is not None:
        level_map = _build_list_level_map(docx_path)
        base_md = _apply_list_indentation(base_md, level_map)

    # 8. Escape block-level HTML tags inside list item text (e.g. <ul>/<ol>)
    #    so renderers don't interpret them as real HTML and break list numbering
    base_md = _escape_block_html_in_lists(base_md)

    parts = [base_md]

    # Tables are already rendered by Docling's export_to_markdown() above.
    # Only append the images/charts appendix — this adds OCR text and chart
    # semantic data which does NOT appear in the base markdown.
    if images:
        parts.append("\n\n---\n## Extracted Image & Chart Data\n")
        for rec in images:
            fname = Path(rec["file"]).name
            parts.append(f"\n### Image {rec['picture_index']}  (`{fname}`)")
            if rec.get("caption"):
                parts.append(f"**Caption:** {rec['caption']}")
            parts.append(f"**Dimensions:** {rec.get('width','?')}×{rec.get('height','?')} px")
            if rec.get("ocr_text"):
                parts.append(f"\n**OCR text:**\n```\n{rec['ocr_text']}\n```")
            sem = rec.get("semantic") or {}
            ct = sem.get("chart_type", "")
            if ct and ct not in ("not_evaluated", "not_a_chart", "unavailable", "error", ""):
                parts.append(f"\n**Chart type:** `{ct}`")
                if sem.get("summary"):
                    parts.append(f"**Summary:** {sem['summary']}")
                for pt in sem.get("series", []):
                    parts.append(f"- {pt.get('label','')}: {pt.get('value','')}")

    return "\n".join(parts)



# ══════════════════════════════════════════════════════════════════════════════
# JSON ENCODER
# ══════════════════════════════════════════════════════════════════════════════

class _SafeEncoder(json.JSONEncoder):
    def default(self, obj):
        try:
            return super().default(obj)
        except TypeError:
            return str(obj)


# ══════════════════════════════════════════════════════════════════════════════
# DOCLING CONVERTER (built once at module import)
# ══════════════════════════════════════════════════════════════════════════════

def _build_docling_converter():
    try:
        from docling.datamodel.base_models import InputFormat
        from docling.datamodel.pipeline_options import PaginatedPipelineOptions
        from docling.document_converter import DocumentConverter, WordFormatOption

        opts = PaginatedPipelineOptions()
        opts.generate_page_images = True
        opts.images_scale = IMAGES_SCALE
        return DocumentConverter(
            format_options={InputFormat.DOCX: WordFormatOption(pipeline_options=opts)}
        )
    except Exception:
        log.warning("PaginatedPipelineOptions unavailable — using default DocumentConverter")
        from docling.document_converter import DocumentConverter
        return DocumentConverter()


_converter = _build_docling_converter()


# ══════════════════════════════════════════════════════════════════════════════
# CORE PIPELINE
# ══════════════════════════════════════════════════════════════════════════════

def _process_document(docx_path: Path) -> dict[str, Any]:
    """Run the full pipeline on a single .docx file."""
    name = docx_path.stem
    out_dir = OUTPUT_DIR / name
    out_dir.mkdir(parents=True, exist_ok=True)

    log.info("Processing: %s", docx_path.name)

    # ── Docling parse ──────────────────────────────────────────────────────────
    try:
        result = _converter.convert(str(docx_path))
    except Exception as exc:
        log.error("Docling failed: %s", exc)
        return {"success": False, "document": docx_path.name, "error": str(exc),
                "markdown": "", "tables_count": 0, "images_count": 0}

    doc = result.document

    # ── DOM JSON ───────────────────────────────────────────────────────────────
    dom_path = out_dir / f"{name}.dom.json"
    try:
        dom_path.write_text(
            json.dumps(doc.export_to_dict(), indent=2, ensure_ascii=False, cls=_SafeEncoder),
            encoding="utf-8",
        )
    except Exception as exc:
        log.warning("DOM export failed: %s", exc)

    # ── Tables + Images ────────────────────────────────────────────────────────
    tables = extract_tables(doc)
    images = extract_images(doc, out_dir / "images")

    # ── Semantic JSON ──────────────────────────────────────────────────────────
    semantic_path = out_dir / f"{name}.semantic.json"
    semantic: dict[str, Any] = {
        "document":          docx_path.name,
        "schema_version":    "1.0",
        "num_pages":         getattr(doc, "num_pages", None),
        "text_blocks_count": len(doc.texts),
        "tables":            tables,
        "images":            images,
    }
    semantic_path.write_text(
        json.dumps(semantic, indent=2, ensure_ascii=False, cls=_SafeEncoder),
        encoding="utf-8",
    )

    # ── Markdown ───────────────────────────────────────────────────────────────
    try:
        base_md = doc.export_to_markdown()
    except Exception as exc:
        base_md = f"*Markdown export failed: {exc}*"

    full_md = render_markdown(base_md, tables, images, docx_path=docx_path)
    md_path = out_dir / f"{name}.md"
    md_path.write_text(full_md, encoding="utf-8")

    log.info("Done: %s  (tables=%d, images=%d)", docx_path.name, len(tables), len(images))

    return {
        "success":            True,
        "document":           docx_path.name,
        "error":              None,
        "markdown":           full_md,
        "dom_json_path":      str(dom_path),
        "semantic_json_path": str(semantic_path),
        "markdown_path":      str(md_path),
        "tables_count":       len(tables),
        "images_count":       len(images),
    }


def _run_pipeline(job_id: str, docx_path: Path) -> None:
    """Worker executed in the thread pool."""
    _jobs[job_id]["status"] = "processing"
    try:
        result = _process_document(docx_path)
        _jobs[job_id].update({
            "status": "done" if result["success"] else "error",
            "result": result,
        })
    except Exception as exc:
        log.exception("Pipeline crashed for job %s", job_id)
        _jobs[job_id].update({
            "status": "error",
            "result": {"success": False, "error": str(exc),
                       "markdown": "", "tables_count": 0, "images_count": 0},
        })
    finally:
        try:
            docx_path.unlink(missing_ok=True)
        except Exception:
            pass


# ══════════════════════════════════════════════════════════════════════════════
# FILE HELPERS
# ══════════════════════════════════════════════════════════════════════════════

import time as _time


def _safe_unlink(path: Path, retries: int = 3, delay: float = 0.5) -> None:
    """
    Delete a file, retrying on Windows PermissionError (WinError 32).
    This happens when doc2docx/Word COM still has the file handle open
    immediately after conversion finishes.
    """
    for attempt in range(retries):
        try:
            path.unlink(missing_ok=True)
            return
        except PermissionError:
            if attempt < retries - 1:
                _time.sleep(delay)
            else:
                log.warning("Could not delete temp file '%s' — still locked. Will be cleaned up later.", path.name)
        except Exception:
            return  # any other error — just move on


# ══════════════════════════════════════════════════════════════════════════════
# LEGACY .doc CONVERSION
# ══════════════════════════════════════════════════════════════════════════════

def _convert_doc_to_docx(doc_path: Path) -> Optional[Path]:
    """
    Convert legacy .doc → .docx.
    Strategy (tried in order):
      1. doc2docx  — uses installed MS Word via COM (Windows only, best quality)
      2. LibreOffice headless — cross-platform, needs soffice on PATH
    Returns the .docx Path on success, None if nothing worked.
    """
    out_path = STAGING_DIR / (doc_path.stem + ".docx")
    STAGING_DIR.mkdir(parents=True, exist_ok=True)

    # ── Strategy 1: doc2docx (MS Word COM, Windows only) ──────────────────────
    try:
        from doc2docx import convert
        convert(str(doc_path), str(out_path))
        if out_path.exists():
            log.info("doc2docx conversion succeeded: %s", out_path.name)
            return out_path
    except ImportError:
        log.info("doc2docx not installed — trying LibreOffice next.")
    except Exception as exc:
        log.warning("doc2docx failed (%s) — trying LibreOffice next.", exc)

    # ── Strategy 2: LibreOffice headless ──────────────────────────────────────
    import subprocess
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "docx",
                 "--outdir", str(STAGING_DIR), str(doc_path)],
                check=True, capture_output=True, timeout=120,
            )
            if out_path.exists():
                log.info("LibreOffice conversion succeeded: %s", out_path.name)
                return out_path
        except Exception as exc:
            log.error("LibreOffice conversion failed: %s", exc)
    else:
        log.warning("LibreOffice not found on PATH.")

    return None


# ══════════════════════════════════════════════════════════════════════════════
# ROUTES
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/upload", response_model=UploadResponse)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
):
    """
    Upload a .doc or .docx file.
    Returns job_id — poll GET /status/{job_id} for progress.
    Output is written to: uploads/<file_stem>/
    """
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in ALLOWED_EXT:
        raise HTTPException(
            status_code=400,
            detail=f"Only .doc and .docx accepted. Got: '{suffix}'",
        )

    job_id = str(uuid.uuid4())
    tmp_path = STAGING_DIR / f"{job_id}{suffix}"

    try:
        with tmp_path.open("wb") as fh:
            shutil.copyfileobj(file.file, fh)
    finally:
        await file.close()

    if suffix == ".doc":
        converted = _convert_doc_to_docx(tmp_path)
        # Use retry-unlink: doc2docx (Word COM) may still hold the file handle
        _safe_unlink(tmp_path)
        if converted is None:
            raise HTTPException(
                status_code=422,
                detail=(
                    "Could not convert legacy .doc file. "
                    "Tried: (1) doc2docx via MS Word COM, (2) LibreOffice headless. "
                    "Please convert to .docx manually and re-upload."
                ),
            )
        tmp_path = converted

    _jobs[job_id] = {
        "status":   "queued",
        "result":   None,
        "filename": file.filename,
    }

    background_tasks.add_task(_executor.submit, _run_pipeline, job_id, tmp_path)

    return UploadResponse(
        job_id=job_id,
        status="queued",
        message=f"'{file.filename}' queued. Poll /status/{job_id}",
    )


@router.get("/status/{job_id}", response_model=JobStatus)
async def get_status(job_id: str):
    """Poll processing status for a job."""
    job = _jobs.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail=f"Job '{job_id}' not found.")

    result: dict = job.get("result") or {}
    md = result.get("markdown", "")

    return JobStatus(
        job_id=job_id,
        status=job["status"],
        document=result.get("document") or job.get("filename"),
        tables_count=result.get("tables_count"),
        images_count=result.get("images_count"),
        markdown_preview=md[:2000] if md else None,
        markdown_path=result.get("markdown_path"),
        semantic_json_path=result.get("semantic_json_path"),
        dom_json_path=result.get("dom_json_path"),
        error=result.get("error"),
    )


@router.get("/download/{job_id}/markdown")
async def download_markdown(job_id: str):
    """Download the generated .md file for a completed job."""
    job = _jobs.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found.")
    if job["status"] != "done":
        raise HTTPException(status_code=409, detail=f"Job not complete (status={job['status']}).")
    md_path = job["result"].get("markdown_path")
    if not md_path or not Path(md_path).exists():
        raise HTTPException(status_code=404, detail="Markdown file not found.")
    return FileResponse(md_path, media_type="text/markdown", filename=Path(md_path).name)


@router.get("/download/{job_id}/semantic")
async def download_semantic(job_id: str):
    """Download the semantic JSON for a completed job."""
    job = _jobs.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found.")
    if job["status"] != "done":
        raise HTTPException(status_code=409, detail="Job not complete.")
    path = job["result"].get("semantic_json_path")
    if not path or not Path(path).exists():
        raise HTTPException(status_code=404, detail="File not found.")
    return FileResponse(path, media_type="application/json", filename=Path(path).name)
