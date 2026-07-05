"""
create_test_docx.py
------------------------------------------------------------------------
Generates a rich .docx test document that exercises every feature of the
Anti-Gravity DOCX -> Markdown pipeline:

  - Multi-section layout headings
  - Ordered and unordered lists
  - Inline text styling (bold, italic, underline)
  - Simple tables (header row + striped rows)
  - Merged-cell table (colspan + rowspan)
  - Embedded bar-chart image (drawn with PIL)
  - Embedded line-chart image
  - Embedded photo-like image (not a chart -- no VLM trigger)
  - Long body text (reading-order check across paragraphs)
  - A footer paragraph

Run:
    python create_test_docx.py
Output:
    f:/Other/anti-gravity/test_docs/AntiGravity_Test.docx
"""

import io
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit(
        "python-docx is not installed. Run:\n"
        "  pip install python-docx\n"
        "or:\n"
        "  pip install python-docx --target .venv\\Lib\\site-packages"
    )

from PIL import Image, ImageDraw, ImageFont
import struct, zlib

OUT_DIR = Path(__file__).parent / "test_docs"
OUT_DIR.mkdir(exist_ok=True)
OUT_PATH = OUT_DIR / "AntiGravity_Test.docx"


# ── helper: shade a table row ────────────────────────────────────────────────
def _shade_row(row, hex_color: str):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)


# ── PIL-based chart image generators ─────────────────────────────────────────

def _pil_to_bytes(img: Image.Image) -> bytes:
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf.read()


def _bar_chart_bytes() -> bytes:
    """Draw a simple bar chart using PIL."""
    W, H = 600, 360
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)

    categories = ["Q1 Claims", "Q2 Claims", "Q3 Claims", "Q4 Claims"]
    values      = [320, 415, 390, 510]
    colors      = ["#4C72B0", "#DD8452", "#55A868", "#C44E52"]
    max_v       = 600
    pad_l, pad_r, pad_t, pad_b = 60, 30, 50, 60
    chart_w = W - pad_l - pad_r
    chart_h = H - pad_t - pad_b

    # Axes
    draw.line([(pad_l, pad_t), (pad_l, H - pad_b)], fill="black", width=2)
    draw.line([(pad_l, H - pad_b), (W - pad_r, H - pad_b)], fill="black", width=2)

    # Y-grid lines
    for tick in [100, 200, 300, 400, 500, 600]:
        y = H - pad_b - int(tick / max_v * chart_h)
        draw.line([(pad_l, y), (W - pad_r, y)], fill="#CCCCCC", width=1)
        draw.text((pad_l - 38, y - 7), str(tick), fill="black")

    # Bars
    bar_w = chart_w // (len(values) * 2)
    for i, (v, c) in enumerate(zip(values, colors)):
        x0 = pad_l + i * (chart_w // len(values)) + bar_w // 2
        y0 = H - pad_b - int(v / max_v * chart_h)
        x1 = x0 + bar_w
        y1 = H - pad_b
        draw.rectangle([x0, y0, x1, y1], fill=c)
        draw.text((x0, y0 - 16), str(v), fill="black")
        draw.text((x0 - 5, H - pad_b + 8), categories[i], fill="black")

    # Title
    draw.text((W // 2 - 120, 10), "Annual Claims by Quarter", fill="black")
    draw.text((8, H // 2 - 10), "Claims", fill="black")
    return _pil_to_bytes(img)


def _line_chart_bytes() -> bytes:
    """Draw a simple dual-line chart using PIL."""
    W, H = 700, 360
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)

    months  = ["Jan","Feb","Mar","Apr","May","Jun",
                "Jul","Aug","Sep","Oct","Nov","Dec"]
    premium = [1200,1250,1180,1310,1400,1380,1420,1350,1500,1460,1530,1600]
    payout  = [800, 750, 820, 900, 950, 870,1000, 980,1050,1020,1100,1150]
    pad_l, pad_r, pad_t, pad_b = 70, 20, 50, 60
    chart_w = W - pad_l - pad_r
    chart_h = H - pad_t - pad_b
    max_v = 1800

    draw.line([(pad_l, pad_t), (pad_l, H - pad_b)], fill="black", width=2)
    draw.line([(pad_l, H - pad_b), (W - pad_r, H - pad_b)], fill="black", width=2)

    for tick in [400, 800, 1200, 1600]:
        y = H - pad_b - int(tick / max_v * chart_h)
        draw.line([(pad_l, y), (W - pad_r, y)], fill="#CCCCCC", width=1)
        draw.text((4, y - 7), str(tick), fill="black")

    step = chart_w // (len(months) - 1)

    def pts(series):
        return [(pad_l + i * step, H - pad_b - int(v / max_v * chart_h))
                for i, v in enumerate(series)]

    prm_pts = pts(premium)
    pay_pts = pts(payout)

    for i in range(len(months) - 1):
        draw.line([prm_pts[i], prm_pts[i+1]], fill="#2196F3", width=3)
        draw.line([pay_pts[i], pay_pts[i+1]], fill="#F44336", width=3)

    for p in prm_pts:
        draw.ellipse([p[0]-4, p[1]-4, p[0]+4, p[1]+4], fill="#2196F3")
    for p in pay_pts:
        draw.ellipse([p[0]-4, p[1]-4, p[0]+4, p[1]+4], fill="#F44336")

    for i, m in enumerate(months):
        draw.text((pad_l + i * step - 8, H - pad_b + 8), m, fill="black")

    draw.text((W // 2 - 160, 10), "Monthly Premium vs Claims (FY 2025)", fill="black")
    draw.rectangle([W-160, 20, W-20, 55], outline="gray")
    draw.line([(W-155, 30), (W-130, 30)], fill="#2196F3", width=3)
    draw.text((W-125, 23), "Premium", fill="black")
    draw.line([(W-155, 45), (W-130, 45)], fill="#F44336", width=3)
    draw.text((W-125, 38), "Payout", fill="black")
    return _pil_to_bytes(img)


def _pie_chart_bytes() -> bytes:
    """Draw a simple pie chart approximation using PIL (rectangles as legend)."""
    W, H = 560, 380
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)

    labels  = ["Auto 35%", "Health 25%", "Property 20%", "Life 12%", "Liability 8%"]
    colors  = ["#FF6384", "#36A2EB", "#FFCE56", "#4BC0C0", "#9966FF"]
    # Draw a crude pie as wedge rectangles in the legend (PIL has no native pie)
    draw.text((W // 2 - 150, 12), "Portfolio Distribution by Line of Business", fill="black")
    cx, cy, r = 200, 210, 140
    import math
    angles = [0.35, 0.25, 0.20, 0.12, 0.08]
    start = -90.0
    for pct, color in zip(angles, colors):
        end = start + pct * 360
        draw.pieslice([cx-r, cy-r, cx+r, cy+r], start=start, end=end, fill=color, outline="white")
        start = end
    # Legend
    for i, (label, color) in enumerate(zip(labels, colors)):
        lx, ly = 380, 100 + i * 35
        draw.rectangle([lx, ly, lx+20, ly+20], fill=color)
        draw.text((lx+26, ly+3), label, fill="black")
    return _pil_to_bytes(img)


def _photo_bytes() -> bytes:
    """A gradient rectangle — simulates a photograph (no chart OCR tokens)."""
    W, H = 480, 320
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    for x in range(W):
        r = int(34  + (x / W) * 80)
        g = int(139 + (x / W) * 60)
        b = int(170 + (x / W) * 85)
        draw.line([(x, 40), (x, H - 30)], fill=(r, g, b))
    draw.rectangle([0, 0, W, 35], fill="#333333")
    draw.text((10, 8), "Site Photograph -- Claim #98421", fill="white")
    draw.rectangle([0, H-28, W, H], fill="#333333")
    draw.text((10, H-22), "GPS: 40.7128 N, 74.0060 W  |  2025-06-15 14:32", fill="white")
    return _pil_to_bytes(img)


# ─────────────────────────────────────────────────────────────────────────────
def build_document() -> None:
    doc = Document()

    # ── Cover / Title ──────────────────────────────────────────────────────────
    title = doc.add_heading("Enterprise Document Parser — Test Sample", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Comprehensive test fixture for the Anti-Gravity DOCX pipeline")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    sub.runs[0].font.size = Pt(12)

    doc.add_paragraph()  # spacer

    # ─── Section 1 ────────────────────────────────────────────────────────────
    doc.add_heading("1. Introduction and Overview", level=1)
    doc.add_paragraph(
        "This document is designed to test every structural element that the "
        "Anti-Gravity pipeline is expected to handle faithfully. The document "
        "deliberately includes complex layouts, merged-cell tables, multiple "
        "embedded chart images, and a plain photograph to ensure that the "
        "chart-heuristic works correctly.\n\n"
        "The pipeline converts this document through the following stages:\n"
        "Docling structural parse → Table extraction → Image OCR (Textract) → "
        "Chart semantics (Bedrock Nova Lite) → Markdown + JSON output."
    )

    # ─── Section 2: Text Formatting ──────────────────────────────────────────
    doc.add_heading("2. Text Formatting & Inline Styles", level=1)
    p = doc.add_paragraph()
    p.add_run("Bold text ").bold = True
    r = p.add_run("Bold text ")
    r.bold = True
    p.add_run("is used for emphasis. ")
    r2 = p.add_run("Italic text ")
    r2.italic = True
    p.add_run("is used for definitions. ")
    r3 = p.add_run("Underlined text ")
    r3.underline = True
    p.add_run("marks important terms.")

    doc.add_paragraph(
        "Reading order across multiple paragraphs is tested here. This paragraph "
        "comes immediately after the formatting demo. The pipeline should preserve "
        "the logical sequence of all text blocks when converting to Markdown."
    )
    doc.add_paragraph(
        "A second body paragraph follows. Docling preserves reading order even "
        "when the source document has multi-column or text-box layouts by "
        "analysing the structural XML rather than visual flow."
    )

    # ─── Ordered list ─────────────────────────────────────────────────────────
    doc.add_heading("2.1 Ordered List", level=2)
    for i, item in enumerate(
        ["Parse DOCX with Docling", "Extract tables → JSON",
         "Extract images → OCR", "Classify charts → Bedrock", "Emit Markdown"], 1
    ):
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{item}")

    # ─── Unordered list ───────────────────────────────────────────────────────
    doc.add_heading("2.2 Bullet List", level=2)
    bullets = [
        "Multi-layout DOCX support",
        "Merged-cell table fidelity",
        "OCR via Amazon Textract (managed, no local binary)",
        "Chart classification via AWS Bedrock Nova Lite",
        "Reading-order preservation",
        "Dual JSON output (DOM + semantic layers)",
    ]
    for item in bullets:
        doc.add_paragraph(item, style="List Bullet")

    # ─── Section 3: Simple Table ──────────────────────────────────────────────
    doc.add_heading("3. Tables", level=1)
    doc.add_heading("3.1 Simple 4×5 Data Table", level=2)
    doc.add_paragraph(
        "The table below represents a sample claims summary report."
    )

    table1 = doc.add_table(rows=5, cols=4)
    table1.style = "Table Grid"
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Claim ID", "Policy Holder", "Type", "Amount (USD)"]
    rows_data = [
        ["CLM-001", "Alice Johnson",   "Auto",     "$4,200"],
        ["CLM-002", "Bob Smith",       "Health",   "$11,500"],
        ["CLM-003", "Carol White",     "Property", "$38,000"],
        ["CLM-004", "David Lee",       "Life",     "$250,000"],
    ]

    hdr_row = table1.rows[0]
    _shade_row(hdr_row, "1F497D")
    for j, h in enumerate(headers):
        cell = hdr_row.cells[j]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, row_data in enumerate(rows_data, start=1):
        row = table1.rows[i]
        if i % 2 == 0:
            _shade_row(row, "DCE6F1")
        for j, val in enumerate(row_data):
            row.cells[j].text = val

    # ─── Merged-cell table ────────────────────────────────────────────────────
    doc.add_heading("3.2 Merged-Cell Table (colspan + rowspan)", level=2)
    doc.add_paragraph(
        "This table tests the pipeline's ability to preserve merged cells "
        "in the raw_cells output. The first two cells in row 1 are merged "
        "horizontally; the first cell in rows 2-3 is merged vertically."
    )

    table2 = doc.add_table(rows=4, cols=3)
    table2.style = "Table Grid"

    # Header — all 3 cells, no merge
    hdr2 = table2.rows[0]
    _shade_row(hdr2, "4F81BD")
    for j, h in enumerate(["Region", "Q1 Loss Ratio", "Q2 Loss Ratio"]):
        cell = hdr2.cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Row 1: merge first two columns (horizontal merge)
    merged = table2.rows[1].cells[0].merge(table2.rows[1].cells[1])
    merged.text = "North America (merged colspan)"
    table2.rows[1].cells[2].text = "67%"

    # Rows 2-3: merge first column vertically (rowspan)
    merged2 = table2.rows[2].cells[0].merge(table2.rows[3].cells[0])
    merged2.text = "Europe (merged rowspan)"
    table2.rows[2].cells[1].text = "58%"
    table2.rows[2].cells[2].text = "61%"
    table2.rows[3].cells[1].text = "54%"
    table2.rows[3].cells[2].text = "59%"

    # ─── Section 4: Charts ────────────────────────────────────────────────────
    doc.add_heading("4. Embedded Images & Charts", level=1)

    # Bar chart
    doc.add_heading("4.1 Bar Chart – Quarterly Claims", level=2)
    doc.add_paragraph(
        "The bar chart below shows the number of claims filed per quarter. "
        "The pipeline should detect this as a chart and call Bedrock for semantic extraction."
    )
    bar_bytes = _bar_chart_bytes()
    if bar_bytes:
        doc.add_picture(io.BytesIO(bar_bytes), width=Inches(5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap = doc.add_paragraph("Figure 1: Quarterly Claims Count")
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.runs[0].italic = True

    # Line chart
    doc.add_heading("4.2 Line Chart – Premium vs Claims", level=2)
    doc.add_paragraph(
        "This multi-series line chart compares monthly premium collected against "
        "claims paid. The Bedrock VLM should identify this as a line chart and "
        "extract the two data series."
    )
    line_bytes = _line_chart_bytes()
    if line_bytes:
        doc.add_picture(io.BytesIO(line_bytes), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap2 = doc.add_paragraph("Figure 2: Premium vs Claims — FY 2025")
    p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap2.runs[0].italic = True

    # Pie chart
    doc.add_heading("4.3 Pie Chart – Portfolio Distribution", level=2)
    doc.add_paragraph(
        "The pie chart illustrates portfolio distribution across lines of "
        "business. The Bedrock VLM should classify this as a pie chart."
    )
    pie_bytes = _pie_chart_bytes()
    if pie_bytes:
        doc.add_picture(io.BytesIO(pie_bytes), width=Inches(4.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap3 = doc.add_paragraph("Figure 3: Portfolio Distribution by LoB")
    p_cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap3.runs[0].italic = True

    # Photo-like image (should NOT trigger VLM chart classification)
    doc.add_heading("4.4 Photograph — Non-Chart Image", level=2)
    doc.add_paragraph(
        "This image simulates a site photograph attached to a property claim. "
        "It should NOT trigger the Bedrock VLM call because it does not look "
        "like a chart (too few OCR tokens or wrong heuristic profile)."
    )
    photo_bytes = _photo_bytes()
    if photo_bytes:
        doc.add_picture(io.BytesIO(photo_bytes), width=Inches(4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap4 = doc.add_paragraph("Figure 4: Site Photograph — Claim #98421")
    p_cap4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap4.runs[0].italic = True

    # ─── Section 5: Multi-layout / Reading order ──────────────────────────────
    doc.add_heading("5. Multi-Layout & Reading Order", level=1)
    doc.add_paragraph(
        "This section tests that the pipeline correctly preserves reading order "
        "even when the document uses multiple heading levels, nested lists, and "
        "sections with varying indentation. The Docling parser operates on the "
        "OOXML structure rather than the rendered page, so reading order is "
        "determined by element position in the XML tree."
    )
    doc.add_heading("5.1 Sub-section Alpha", level=2)
    doc.add_paragraph(
        "Alpha content paragraph. Lorem ipsum dolor sit amet, consectetur "
        "adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore "
        "magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation."
    )
    doc.add_heading("5.2 Sub-section Beta", level=2)
    doc.add_paragraph(
        "Beta content paragraph. Duis aute irure dolor in reprehenderit in "
        "voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur "
        "sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt "
        "mollit anim id est laborum."
    )
    doc.add_heading("5.3 Nested List", level=2)
    for top in ["Coverage Exclusions", "Claim Procedures", "Premium Adjustments"]:
        p = doc.add_paragraph(top, style="List Bullet")
        for sub in ["Sub-item A", "Sub-item B"]:
            p2 = doc.add_paragraph(f"    {sub}", style="List Bullet 2")

    # ─── Section 6: Summary ───────────────────────────────────────────────────
    doc.add_heading("6. Summary & Conclusions", level=1)
    doc.add_paragraph(
        "This test document covers the full breadth of features supported by "
        "the Anti-Gravity pipeline:\n\n"
        "• Structural text (headings H1–H3, paragraphs)\n"
        "• Inline formatting (bold, italic, underline)\n"
        "• Ordered and unordered lists\n"
        "• Simple data tables with header rows\n"
        "• Merged-cell tables (colspan and rowspan)\n"
        "• Embedded chart images (bar, line, pie)\n"
        "• Non-chart images (photo)\n"
        "• Reading-order across multiple sections\n\n"
        "After processing, the output directory should contain:\n"
        "  AntiGravity_Test/AntiGravity_Test.md\n"
        "  AntiGravity_Test/AntiGravity_Test.dom.json\n"
        "  AntiGravity_Test/AntiGravity_Test.semantic.json\n"
        "  AntiGravity_Test/images/picture_0.png  (bar chart)\n"
        "  AntiGravity_Test/images/picture_1.png  (line chart)\n"
        "  AntiGravity_Test/images/picture_2.png  (pie chart)\n"
        "  AntiGravity_Test/images/picture_3.png  (photo — no VLM)"
    )

    # ─── Footer note ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        "Generated by create_test_docx.py  |  Anti-Gravity Pipeline Test Suite  |  v1.0"
    )
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(str(OUT_PATH))
    print("[OK] Test document written to: " + str(OUT_PATH))


if __name__ == "__main__":
    build_document()
