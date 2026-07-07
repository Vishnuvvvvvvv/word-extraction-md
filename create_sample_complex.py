"""
create_sample_complex.py
────────────────────────
Generates a rich test .docx with:
  - Multiple heading levels
  - Simple bullet list
  - Nested (multi-level) bullet list
  - Simple data table
  - Merged-cell table (colspan + rowspan)
  - Table with a nested mini-table inside a cell

Run:
    python create_sample_complex.py
Output:
    test_docs/Complex_Test.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT_DIR = Path("test_docs")
OUT_DIR.mkdir(exist_ok=True)
OUT_FILE = OUT_DIR / "Complex_Test.docx"

doc = Document()


# ── Helpers ───────────────────────────────────────────────────────────────────

def heading(text, level=1):
    doc.add_heading(text, level=level)

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def _set_list_level(paragraph, level: int, num_id: int):
    """
    Apply proper OOXML list numbering (numId + ilvl) so Docling
    reads the indentation level correctly.
    level 0 = top-level, level 1 = one indent, etc.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numId_el = OxmlElement('w:numId')
    numId_el.set(qn('w:val'), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId_el)
    pPr.append(numPr)

def bullet(text, level=0):
    """Add a bullet list item with proper OOXML ilvl so Docling detects nesting."""
    p = doc.add_paragraph(text, style="List Bullet")
    _set_list_level(p, level, num_id=1)
    return p

def numbered(text, level=0):
    """Add a numbered list item with proper OOXML ilvl so Docling detects nesting."""
    p = doc.add_paragraph(text, style="List Number")
    _set_list_level(p, level, num_id=2)
    return p

def shade_cell(cell, hex_color="D9E1F2"):
    """Apply background colour to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def bold_cell(cell, text):
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.bold = True

def merge_h(row, start_col, end_col):
    """Merge cells horizontally in a row."""
    row.cells[start_col].merge(row.cells[end_col])

def merge_v(table, col, start_row, end_row):
    """Merge cells vertically in a column."""
    table.cell(start_row, col).merge(table.cell(end_row, col))

def add_nested_table(cell, data):
    """
    Add a small table inside a cell.
    data = list of rows, each row is a list of strings.
    """
    nested = cell.add_table(rows=len(data), cols=len(data[0]))
    nested.style = "Table Grid"
    for r, row_data in enumerate(data):
        for c, text in enumerate(row_data):
            nested.cell(r, c).text = text
    return nested


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — Introduction & Headings
# ══════════════════════════════════════════════════════════════════════════════

heading("Complex Document Test", level=1)
para("This document is designed to test the Anti-Gravity extraction pipeline "
     "against a wide variety of Word features: bullet lists, nested lists, "
     "simple tables, merged-cell tables, and nested tables inside cells.")

heading("1. Bullet Lists", level=2)

# Simple bullets
heading("1.1 Simple bullet list", level=4)
for item in ["Alpha item", "Beta item", "Gamma item", "Delta item"]:
    bullet(item, level=0)

doc.add_paragraph()

# Nested bullets
heading("1.2 Nested bullet list", level=4)
bullet("Frontend Technologies", level=0)
bullet("React.js", level=1)
bullet("Vite", level=1)
bullet("Vanilla CSS", level=1)
bullet("Backend Technologies", level=0)
bullet("FastAPI", level=1)
bullet("Docling", level=1)
bullet("Amazon Textract", level=1)
bullet("AWS Bedrock Nova Lite", level=1)
bullet("Infrastructure", level=0)
bullet("Docker", level=1)
bullet("Kubernetes", level=1)
bullet("AWS ECS", level=1)

doc.add_paragraph()

# Numbered list with nested sub-items
heading("2. Numbered & Nested Lists", level=2)
para("Processing pipeline steps:", bold=True)
numbered("Upload DOCX via API", level=0)
numbered("Validate file extension", level=1)
numbered("Save to staging folder", level=1)
numbered("Docling structural parse", level=0)
numbered("Extract text blocks", level=1)
numbered("Extract tables", level=1)
numbered("Extract embedded images", level=1)
numbered("OCR + Chart Analysis", level=0)
numbered("Amazon Textract → raw text", level=1)
numbered("Bedrock Nova Lite → chart semantics", level=1)
numbered("Assemble Markdown + JSON output", level=0)

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — Simple Data Table
# ══════════════════════════════════════════════════════════════════════════════

heading("3. Simple Data Table", level=2)
para("A standard 5-row claims summary table with a header row.")

t1 = doc.add_table(rows=6, cols=4)
t1.style = "Table Grid"

headers = ["Claim ID", "Policy Holder", "Type", "Amount (USD)"]
for i, h in enumerate(headers):
    bold_cell(t1.rows[0].cells[i], h)
    shade_cell(t1.rows[0].cells[i], "4472C4")
    t1.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

data = [
    ["CLM-001", "Alice Johnson",  "Auto",     "$4,200"],
    ["CLM-002", "Bob Smith",      "Health",   "$11,500"],
    ["CLM-003", "Carol White",    "Property", "$38,000"],
    ["CLM-004", "David Lee",      "Life",     "$250,000"],
    ["CLM-005", "Eva Martinez",   "Travel",   "$3,750"],
]
for r, row_data in enumerate(data, start=1):
    for c, text in enumerate(row_data):
        t1.rows[r].cells[c].text = text

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — Merged-Cell Table (colspan + rowspan)
# ══════════════════════════════════════════════════════════════════════════════

heading("4. Merged-Cell Table", level=2)
para("This table uses both colspan (horizontal merge) and rowspan (vertical merge).")

# 4 rows × 4 cols
t2 = doc.add_table(rows=5, cols=4)
t2.style = "Table Grid"

# Header row — first two cells merged (colspan)
for i, h in enumerate(["Region & Quarter", "Region & Quarter", "Q1 Loss %", "Q2 Loss %"]):
    bold_cell(t2.rows[0].cells[i], h)
    shade_cell(t2.rows[0].cells[i], "ED7D31")
merge_h(t2.rows[0], 0, 1)

# Row 1+2 — first col merged (rowspan) — North America
t2.rows[1].cells[0].text = "North America"
t2.rows[1].cells[1].text = "East Coast"
t2.rows[1].cells[2].text = "62%"
t2.rows[1].cells[3].text = "67%"

t2.rows[2].cells[0].text = "North America"
t2.rows[2].cells[1].text = "West Coast"
t2.rows[2].cells[2].text = "58%"
t2.rows[2].cells[3].text = "61%"

merge_v(t2, 0, 1, 2)

# Row 3+4 — first col merged (rowspan) — Europe
t2.rows[3].cells[0].text = "Europe"
t2.rows[3].cells[1].text = "UK"
t2.rows[3].cells[2].text = "55%"
t2.rows[3].cells[3].text = "59%"

t2.rows[4].cells[0].text = "Europe"
t2.rows[4].cells[1].text = "Germany"
t2.rows[4].cells[2].text = "51%"
t2.rows[4].cells[3].text = "54%"

merge_v(t2, 0, 3, 4)

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — Table with Nested Table Inside a Cell
# ══════════════════════════════════════════════════════════════════════════════

heading("5. Table with Nested Table in a Cell", level=2)
para("The 'Coverage Details' cell below contains a mini nested table "
     "listing coverage tiers and their limits.")

t3 = doc.add_table(rows=3, cols=3)
t3.style = "Table Grid"

# Header row
for i, h in enumerate(["Product", "Coverage Details", "Premium (Monthly)"]):
    bold_cell(t3.rows[0].cells[i], h)
    shade_cell(t3.rows[0].cells[i], "70AD47")
    t3.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# Row 1 — simple data
t3.rows[1].cells[0].text = "Basic Plan"
t3.rows[1].cells[2].text = "$150"
# Insert nested table into middle cell
add_nested_table(t3.rows[1].cells[1], [
    ["Tier",     "Limit"],
    ["Hospital", "$50,000"],
    ["Surgery",  "$20,000"],
    ["OPD",      "$5,000"],
])

# Row 2 — another nested table
t3.rows[2].cells[0].text = "Premium Plan"
t3.rows[2].cells[2].text = "$380"
add_nested_table(t3.rows[2].cells[1], [
    ["Tier",         "Limit"],
    ["Hospital",     "$2,00,000"],
    ["Critical Ill", "$1,00,000"],
    ["Dental",       "$15,000"],
    ["Vision",       "$8,000"],
])

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — Mixed Content Table (bullets inside cells)
# ══════════════════════════════════════════════════════════════════════════════

heading("6. Table with Bullet Points Inside Cells", level=2)
para("Each cell in the 'Features' column contains a bullet list.")

t4 = doc.add_table(rows=3, cols=2)
t4.style = "Table Grid"

bold_cell(t4.rows[0].cells[0], "Plan")
bold_cell(t4.rows[0].cells[1], "Key Features")
shade_cell(t4.rows[0].cells[0], "7030A0")
shade_cell(t4.rows[0].cells[1], "7030A0")
for cell in t4.rows[0].cells:
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

t4.rows[1].cells[0].text = "Silver"
cell_silver = t4.rows[1].cells[1]
cell_silver.paragraphs[0].clear()
for feat in ["Hospital cover", "OPD included", "No waiting period"]:
    p = cell_silver.add_paragraph(feat, style="List Bullet")

t4.rows[2].cells[0].text = "Gold"
cell_gold = t4.rows[2].cells[1]
cell_gold.paragraphs[0].clear()
for feat in ["All Silver features", "Critical illness cover", "International emergency", "Annual health checkup"]:
    p = cell_gold.add_paragraph(feat, style="List Bullet")

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — Summary
# ══════════════════════════════════════════════════════════════════════════════

heading("7. Summary", level=2)
para("This document was generated to test every structural element "
     "the Anti-Gravity pipeline must handle:")

bullet("Simple bullet list (top-level)", level=0)
bullet("Nested bullet list (two levels)", level=0)
bullet("Numbered list with sub-items", level=0)
bullet("Simple data table (5 rows × 4 cols)", level=0)
bullet("Merged-cell table (colspan + rowspan)", level=0)
bullet("Table with nested mini-table inside a cell", level=0)
bullet("Table with bullet list items inside cells", level=0)

doc.add_paragraph()
para("Generated by create_sample_complex.py  |  Anti-Gravity Pipeline Test Suite", italic=True)


# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT_FILE)
print(f"[OK] Written to: {OUT_FILE.resolve()}")
